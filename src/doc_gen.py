# src/doc_gen.py

import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

class DocumentGenerator:
    def __init__(self):
        pass

    def convert_markdown_to_docx(self, markdown_text: str, output_path: str):
        """
        将 Markdown 转换为 Word，支持：
        1. 自动去除行首空格（解决解析失败问题）
        2. 正确渲染 **加粗** 文字
        3. 智能寻找图片路径
        """
        doc = Document()
        self._set_global_style(doc)

        # 获取文档所在的基准目录 (例如 ./output)
        base_dir = os.path.dirname(os.path.abspath(output_path))
        print(f"📂 文档基准路径: {base_dir}")

        lines = markdown_text.split('\n')
        
        for line in lines:
            # 关键修复 1: 去除首尾空格，防止 "  ## 标题" 识别失败
            stripped_line = line.strip()
            
            if not stripped_line:
                continue
            
            # --- 标题处理 ---
            if stripped_line.startswith('# '):
                self._add_heading(doc, stripped_line[2:], level=1)
            elif stripped_line.startswith('## '):
                self._add_heading(doc, stripped_line[3:], level=2)
            elif stripped_line.startswith('### '):
                self._add_heading(doc, stripped_line[4:], level=3)
            elif stripped_line.startswith('#### '):
                self._add_heading(doc, stripped_line[5:], level=4)
            
            # --- 图片处理 ---
            # 兼容 ![alt](path) 和 HTML 格式的图片标签
            elif stripped_line.startswith('![') and '](' in stripped_line:
                self._add_image(doc, stripped_line, base_dir)
            
            # --- 分隔线 ---
            elif stripped_line == '---' or stripped_line == '***':
                doc.add_page_break()
                
            # --- 引用 ---
            elif stripped_line.startswith('> '):
                p = doc.add_paragraph()
                p.style = 'Intense Quote'
                self._render_rich_text(p, stripped_line[2:])
                
            # --- 普通段落 ---
            else:
                p = doc.add_paragraph()
                # 关键修复 2: 调用富文本渲染，处理 **加粗**
                self._render_rich_text(p, stripped_line)

        # 保存
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        try:
            doc.save(output_path)
            print(f"✅ Word 文档生成成功: {output_path}")
        except Exception as e:
            print(f"❌ 无法保存文件 (可能文件被占用): {e}")

    def _set_global_style(self, doc):
        """设置中西文混合字体"""
        try:
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)
            style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        except:
            pass

    def _add_heading(self, doc, text, level):
        # 同样支持标题中的加粗渲染
        heading = doc.add_heading(level=level)
        self._render_rich_text(heading, text)
        
        # 样式调整
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.name = '黑体' # 标题倾向于黑体
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    def _render_rich_text(self, paragraph, text):
        """
        解析 Markdown 的 **加粗** 语法并应用到 Word 段落
        """
        # 正则拆分: (非加粗部分, 加粗部分, 非加粗部分...)
        # pattern matching **text**
        parts = re.split(r'(\*\*.*?\*\*)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # 去掉 ** 并加粗
                content = part[2:-2]
                run = paragraph.add_run(content)
                run.bold = True
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            else:
                if part: # 防止空字符串
                    run = paragraph.add_run(part)
                    run.font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def _add_image(self, doc, line, base_dir):
        """
        更加智能的路径查找逻辑
        """
        # 提取路径
        match = re.search(r'\!\[.*?\]\((.*?)\)', line)
        if not match:
            return
            
        raw_path = match.group(1).strip()
        
        # 路径清理：有些模型会输出 assets/img.jpg "Title"
        if " " in raw_path and raw_path.lower().endswith(('jpg', 'png', 'jpeg"')):
             raw_path = raw_path.split(" ")[0]
        
        # 移除可能存在的引号
        raw_path = raw_path.strip('"').strip("'")

        # --- 多级路径探测 ---
        candidates = [
            raw_path,                                      # 1. 绝对路径或相对于运行目录
            os.path.join(base_dir, raw_path),              # 2. 相对于 docx 输出目录
            os.path.join(base_dir, os.path.basename(raw_path)), # 3. 甚至直接在 assets 平级找
            os.path.abspath(raw_path)                      # 4. 绝对路径
        ]
        
        # 如果路径以 ./ 开头，尝试去掉
        if raw_path.startswith("./"):
            clean_path = raw_path[2:]
            candidates.append(os.path.join(base_dir, clean_path))

        final_path = None
        for p in candidates:
            # 统一分隔符
            p = p.replace("/", os.sep).replace("\\", os.sep)
            if os.path.exists(p) and os.path.isfile(p):
                final_path = p
                break
        
        if final_path:
            try:
                # 插入图片
                doc.add_picture(final_path, width=Inches(6.0))
                last_p = doc.paragraphs[-1]
                last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                print(f"  🖼️  图片插入成功: {os.path.basename(final_path)}")
            except Exception as e:
                print(f"  ❌ 图片文件损坏或格式不支持: {final_path}")
                p = doc.add_paragraph(f"[图片格式错误: {os.path.basename(final_path)}]")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            print(f"  ⚠️ 图片未找到 (Markdown路径: {raw_path})")
            print(f"     尝试过查找: {candidates}")
            # 在文档里留个红色的提示
            p = doc.add_paragraph(f"[图片丢失: {raw_path}]")
            p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
